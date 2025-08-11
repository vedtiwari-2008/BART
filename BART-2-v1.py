#!/usr/bin/env python3
"""
BART (Bigis Automated Rank Tracer) - Complete GUI Application
A modern, professional Google search ranking tracker with Bigis Technology branding.
Single-file application with configuration window and main tracking interface.
"""

import customtkinter as ctk
import tkinter as tk
from tkinter import filedialog, messagebox
import threading
import os
import sys
from datetime import datetime
import logging
import undetected_chromedriver as uc
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from urllib.parse import urlparse
import traceback
import re
import time
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('bart_gui.log'),
        logging.StreamHandler()
    ]
)

# Set customtkinter appearance
ctk.set_appearance_mode("dark")
ctk.set_default_color_theme("blue")

# Bigis Technology Color Scheme
BIGIS_COLORS = {
    'primary': '#003366',      # Deep blue
    'secondary': '#0066cc',    # Medium blue
    'accent': '#ff6b35',       # Orange accent
    'success': '#28a745',      # Green
    'warning': '#ffc107',      # Yellow
    'danger': '#dc3545',       # Red
    'light': '#f8f9fa',        # Light gray
    'dark': '#1a1a1a',         # Dark background
    'gray': '#6c757d',         # Medium gray
    'white': '#ffffff'
}

# Font color options for Word document
FONT_COLORS = {
    'Black': RGBColor(0, 0, 0),
    'Blue': RGBColor(0, 51, 102),
    'Red': RGBColor(220, 53, 69),
    'Green': RGBColor(40, 167, 69),
    'Dark Gray': RGBColor(52, 58, 64),
    'Purple': RGBColor(111, 66, 193),
    'Orange': RGBColor(255, 107, 53)
}

# ==================== UTILITY FUNCTIONS ====================

def clean_domain(url):
    """Clean and normalize domain"""
    try:
        if not url or not isinstance(url, str):
            return ""
        
        url = url.strip().lower()
        parsed = urlparse(url)
        domain = parsed.netloc
        
        if not domain:
            domain_match = re.search(r'(?:https?://)?(?:www\.)?([^/\s?]+)', url)
            if domain_match:
                domain = domain_match.group(1)
        
        domain = re.sub(r'^www\.', '', domain)
        domain = re.sub(r'^m\.', '', domain)
        domain = re.sub(r'^mobile\.', '', domain)
        domain = re.sub(r':\d+$', '', domain)
        
        domain = domain.strip().lower()
        
        if '.' not in domain or len(domain) < 3:
            return ""
            
        return domain
        
    except Exception as e:
        logging.warning(f"Error cleaning domain '{url}': {str(e)}")
        return ""

def is_target_match(found_domain, target_domain):
    """Check if domains match"""
    if not found_domain or not target_domain:
        return False
    
    found_clean = clean_domain(found_domain)
    target_clean = clean_domain(target_domain)
    
    if not found_clean or not target_clean:
        return False
    
    # Exact match
    if found_clean == target_clean:
        return True
    
    # Subdomain check
    if found_clean.endswith('.' + target_clean) or target_clean.endswith('.' + found_clean):
        return True
    
    # Main domain comparison
    target_parts = target_clean.split('.')
    found_parts = found_clean.split('.')
    
    if len(target_parts) >= 2 and len(found_parts) >= 2:
        target_main = '.'.join(target_parts[-2:])
        found_main = '.'.join(found_parts[-2:])
        if target_main == found_main:
            return True
    
    return False

def wait_for_search_ready(driver, timeout=45):
    """Wait for search box to be ready"""
    start_time = time.time()
    
    while time.time() - start_time < timeout:
        try:
            search_box = driver.find_element(By.NAME, "q")
            if search_box.is_enabled() and search_box.is_displayed():
                try:
                    search_box.click()
                    return search_box
                except:
                    pass
        except:
            pass
        
        # Check for captcha
        if "recaptcha" in driver.page_source.lower():
            print("🤖 Captcha detected - solve manually")
        
        time.sleep(0.5)
    
    try:
        return driver.find_element(By.NAME, "q")
    except:
        raise Exception("Search box not available after waiting")

def get_top_10_organic_results(driver):
    """Get exactly the top 10 main organic search results"""
    main_selectors = [
        "div.g div.yuRUbf a[href]:not([href*='google.com'])",
        "div.tF2Cxc div.yuRUbf a[href]:not([href*='google.com'])",
    ]
    
    for selector in main_selectors:
        try:
            elements = driver.find_elements(By.CSS_SELECTOR, selector)
            
            valid_results = []
            for element in elements:
                href = element.get_attribute('href')
                if href and is_main_organic_result(href):
                    try:
                        parent_container = element.find_element(By.XPATH, "./ancestor::div[contains(@class, 'g') or contains(@class, 'tF2Cxc')]")
                        parent_html = parent_container.get_attribute('outerHTML').lower()
                        
                        exclude_sections = [
                            'ads-fr', 'commercial', 'sponsored', 'ad_cclk',
                            'people also ask', 'related questions', 'accordion',
                            'related searches', 'knowledge panel', 'knowledge-panel',
                            'kno-kp', 'kp-', 'g-blk', 'mnr-c', 'UDZeY', 'akr-n',
                            'related-question', 'accordion-toggle'
                        ]
                        
                        if not any(exclude in parent_html for exclude in exclude_sections):
                            try:
                                title_elem = parent_container.find_element(By.CSS_SELECTOR, "h3")
                                desc_elem = parent_container.find_element(By.CSS_SELECTOR, "div[data-sncf], .VwiC3b, .s3v9rd, [data-content-feature]")
                                if title_elem and desc_elem:
                                    valid_results.append(element)
                            except:
                                continue
                    except:
                        continue
            
            if valid_results:
                return valid_results[:10]
                
        except Exception as e:
            continue
    
    return []

def is_main_organic_result(url):
    """Check if URL is a main organic result"""
    if not url:
        return False
    
    exclude_patterns = [
        'google.com', 'googleusercontent.com', 'youtube.com/redirect',
        'accounts.google', 'support.google', 'policies.google',
        'webcache.googleusercontent', 'translate.google', 'maps.google',
        'shopping.google', 'images.google', 'news.google',
        'javascript:', 'mailto:', '/search?', '/preferences?',
        'tbm=isch', 'tbm=vid', 'tbm=nws', 'googleadservices',
        'googlesyndication', '/aclk?', '/url?q=', 'googleads',
    ]
    
    url_lower = url.lower()
    for pattern in exclude_patterns:
        if pattern in url_lower:
            return False
    
    if not (url_lower.startswith('http://') or url_lower.startswith('https://')):
        return False
    
    return True

def get_title_safe(driver, link):
    """Get page title safely"""
    try:
        title_selectors = [
            "./ancestor::div[contains(@class, 'g')]//h3",
            "./ancestor::div[contains(@class, 'tF2Cxc')]//h3",
            "./ancestor::div[contains(@class, 'yuRUbf')]//h3",
            ".//h3",
            "./parent::*//h3"
        ]
        
        for selector in title_selectors:
            try:
                title_elem = link.find_element(By.XPATH, selector)
                title = title_elem.text.strip()
                if title and len(title) > 3:
                    return title
            except:
                continue
        
        return "No title found"
        
    except:
        return "Title error"

# ==================== RANK TRACKER CLASS ====================

class RankTracker:
    """Enhanced rank tracker with GUI integration and Word document generation"""
    
    def __init__(self, keyword, target_domain, max_pages, config, log_callback=None, status_callback=None):
        self.keyword = keyword
        self.target_domain = target_domain
        self.max_pages = max_pages
        self.config = config
        self.log_callback = log_callback or (lambda msg: print(msg))
        self.status_callback = status_callback or (lambda msg: None)
        
        self.driver = None
        self.found_result = None
    
    def log(self, message):
        """Log message using callback"""
        self.log_callback(message)
        logging.info(message)
    
    def update_status(self, status):
        """Update status using callback"""
        self.status_callback(status)
    
    def track_ranking(self):
        """Main tracking method"""
        try:
            self.log(f"🎯 Searching: '{self.keyword}' -> {self.target_domain}")
            self.log(f"📄 Max Pages: {self.max_pages}")
            self.log("-" * 50)
            
            self.update_status("Setting up Chrome browser...")
            
            # Chrome options for visible browser (for CAPTCHA solving)
            options = uc.ChromeOptions()
            options.add_argument("--no-sandbox")
            options.add_argument("--disable-dev-shm-usage")
            options.add_argument("--disable-gpu")
            options.add_argument("--disable-extensions")
            options.add_argument("--disable-plugins")
            options.add_argument("--disable-background-timer-throttling")
            options.add_argument("--disable-backgrounding-occluded-windows")
            options.add_argument("--disable-renderer-backgrounding")
            options.add_argument("--disable-features=TranslateUI")
            options.add_argument("--no-first-run")
            options.add_argument("--no-default-browser-check")
            # Keep browser visible for CAPTCHA solving
            
            self.driver = uc.Chrome(options=options)
            self.driver.get("https://www.google.com")
            
            self.update_status("Waiting for search to be ready...")
            self.log("🌐 Opening Google...")
            
            search_box = wait_for_search_ready(self.driver)
            
            self.update_status("Performing search...")
            self.log("🔍 Entering search query...")
            
            try:
                search_box.clear()
            except:
                try:
                    search_box.click()
                    time.sleep(1)
                    search_box.send_keys(Keys.CONTROL + "a")
                    time.sleep(0.5)
                except:
                    pass
            
            search_box.send_keys(self.keyword)
            time.sleep(0.5)
            search_box.send_keys(Keys.RETURN)
            time.sleep(2)
            
            target_clean = clean_domain(self.target_domain)
            overall_position = 0
            
            # Search through pages
            for page_num in range(1, self.max_pages + 1):
                try:
                    self.update_status(f"Scanning page {page_num} of {self.max_pages}...")
                    self.log(f"📄 Scanning page {page_num}...")
                    
                    WebDriverWait(self.driver, 10).until(
                        EC.presence_of_element_located((By.CSS_SELECTOR, "div.g, div.tF2Cxc"))
                    )
                    
                    links = get_top_10_organic_results(self.driver)
                    self.log(f"🔍 Found {len(links)} organic results on page {page_num}")
                    
                    page_start_position = (page_num - 1) * 10
                    
                    for i, link in enumerate(links):
                        try:
                            url = link.get_attribute('href')
                            title = get_title_safe(self.driver, link)
                            domain = clean_domain(url)
                            position = page_start_position + i + 1
                            overall_position = position
                            
                            self.log(f"  #{position}: {domain} - {title[:50]}...")
                            
                            if is_target_match(domain, target_clean):
                                self.log(f"🎯 FOUND! Target domain at position #{position}")
                                self.log(f"   URL: {url}")
                                self.log(f"   Title: {title}")
                                
                                result = {
                                    'keyword': self.keyword,
                                    'target_domain': self.target_domain,
                                    'position': position,
                                    'page': page_num,
                                    'url': url,
                                    'title': title,
                                    'found': True
                                }
                                
                                self.driver.quit()
                                return result
                                
                        except Exception as e:
                            self.log(f"⚠️ Error processing result {i+1}: {str(e)}")
                            continue
                    
                    # Navigate to next page if not found
                    if page_num < self.max_pages:
                        try:
                            next_button = self.driver.find_element(By.ID, "pnnext")
                            if next_button.is_enabled():
                                next_button.click()
                                time.sleep(3)
                            else:
                                self.log("⚠️ Next button not available")
                                break
                        except:
                            self.log("⚠️ Could not find next page button")
                            break
                    
                except Exception as e:
                    self.log(f"❌ Error on page {page_num}: {str(e)}")
                    continue
            
            # Not found
            self.log(f"❌ Target domain not found in top {overall_position} results")
            result = {
                'keyword': self.keyword,
                'target_domain': self.target_domain,
                'position': 0,
                'page': 0,
                'url': '',
                'title': '',
                'found': False
            }
            
            self.driver.quit()
            return result
            
        except Exception as e:
            self.log(f"❌ Fatal error: {str(e)}")
            if self.driver:
                try:
                    self.driver.quit()
                except:
                    pass
            
            return {
                'keyword': self.keyword,
                'target_domain': self.target_domain,
                'position': 0,
                'page': 0,
                'url': '',
                'title': '',
                'found': False,
                'error': str(e)
            }
    
    def create_word_document(self, result):
        """Create or append to Word document with ranking result"""
        try:
            file_path = os.path.join(self.config['save_location'], f"{self.config['filename']}.docx")
            
            # Check if file already exists
            if os.path.exists(file_path):
                # Open existing document
                doc = Document(file_path)
                self.log(f"📄 Appending to existing Word document...")
                
                # Add new result with spacing
                doc.add_paragraph()  # Add spacing before new result
                
                # Skip header creation since it already exists
                create_header = False
            else:
                # Create new document
                doc = Document()
                self.log(f"📄 Creating new Word document...")
                create_header = True
            
            # Set document styling
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Calibri'
            font.size = Pt(self.config['font_size'])
            font.color.rgb = self.config['font_color']
            
            # Only add header for new documents
            if create_header:
                # Add header with Bigis Technology branding
                header = doc.add_heading('BART Ranking Report', 0)
                header.alignment = WD_ALIGN_PARAGRAPH.CENTER
                header_run = header.runs[0]
                header_run.font.color.rgb = RGBColor(0, 51, 102)  # Bigis blue
                
                subtitle = doc.add_paragraph()
                subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
                subtitle_run = subtitle.add_run('Bigis Technology - Professional SEO Analytics')
                subtitle_run.font.size = Pt(12)
                subtitle_run.font.color.rgb = RGBColor(102, 102, 102)
                subtitle_run.italic = True
                
                doc.add_paragraph()  # Spacing
                
                # Add timestamp
                timestamp = doc.add_paragraph()
                timestamp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                time_run = timestamp.add_run(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
                time_run.font.size = Pt(10)
                time_run.font.color.rgb = RGBColor(102, 102, 102)
                
                doc.add_paragraph()  # Spacing
            
            # Add result (this works for both new and existing documents)
            result_para = doc.add_paragraph()
            result_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            if result['found']:
                result_text = f"{result['keyword']} = Page {result['page']}"
            else:
                result_text = f"{result['keyword']} = Not Found"
            
            result_run = result_para.add_run(result_text)
            result_run.font.name = 'Calibri'
            result_run.font.size = Pt(self.config['font_size'])
            result_run.font.color.rgb = self.config['font_color']
            result_run.bold = True
            
            # Save document (overwrites if existing, creates if new)
            doc.save(file_path)
            
            if os.path.exists(file_path):
                self.log(f"📄 Word document updated: {file_path}")
            else:
                self.log(f"📄 Word document created: {file_path}")
            
            return file_path
            
        except Exception as e:
            self.log(f"❌ Error creating Word document: {str(e)}")
            raise e

# ==================== CONFIGURATION WINDOW ====================

class ConfigurationWindow:
    """Configuration window for BART settings"""
    
    def __init__(self, on_complete_callback):
        self.on_complete_callback = on_complete_callback
        self.window = None
        self.config_data = {}
        
        self.create_window()
    
    def create_window(self):
        """Create the configuration window"""
        self.window = ctk.CTk()
        self.window.title("BART Configuration - Bigis Technology")
        self.window.geometry("600x700")
        self.window.resizable(False, False)
        
        # Handle window close event
        self.window.protocol("WM_DELETE_WINDOW", self.on_closing)
        
        # Configure grid
        self.window.grid_columnconfigure(0, weight=1)
        
        # Main frame
        main_frame = ctk.CTkFrame(self.window)
        main_frame.grid(row=0, column=0, sticky="nsew", padx=20, pady=20)
        main_frame.grid_columnconfigure(0, weight=1)
        
        # Header
        header_frame = ctk.CTkFrame(main_frame, fg_color=BIGIS_COLORS['primary'])
        header_frame.grid(row=0, column=0, sticky="ew", padx=0, pady=0)
        header_frame.grid_columnconfigure(0, weight=1)
        
        # Logo and title
        logo_label = ctk.CTkLabel(
            header_frame, 
            text="📊 BART", 
            font=ctk.CTkFont(size=32, weight="bold"),
            text_color=BIGIS_COLORS['white']
        )
        logo_label.grid(row=0, column=0, pady=(20, 5))
        
        subtitle_label = ctk.CTkLabel(
            header_frame,
            text="Bigis Automated Rank Tracer",
            font=ctk.CTkFont(size=16),
            text_color=BIGIS_COLORS['light']
        )
        subtitle_label.grid(row=1, column=0, pady=(0, 10))
        
        brand_label = ctk.CTkLabel(
            header_frame,
            text="Bigis Technology - Professional SEO Analytics",
            font=ctk.CTkFont(size=12, slant="italic"),
            text_color=BIGIS_COLORS['accent']
        )
        brand_label.grid(row=2, column=0, pady=(0, 20))
        
        # Configuration form
        form_frame = ctk.CTkFrame(main_frame, fg_color="transparent")
        form_frame.grid(row=1, column=0, sticky="ew", padx=20, pady=20)
        form_frame.grid_columnconfigure(1, weight=1)
        
        # File name
        ctk.CTkLabel(form_frame, text="Word Report File Name:", font=ctk.CTkFont(weight="bold")).grid(
            row=0, column=0, sticky="w", pady=(0, 5))
        self.filename_entry = ctk.CTkEntry(form_frame, placeholder_text="Enter filename (without .docx)")
        self.filename_entry.grid(row=0, column=1, sticky="ew", padx=(10, 0), pady=(0, 5))
        self.filename_entry.insert(0, "BART_Report")
        
        # Font size
        ctk.CTkLabel(form_frame, text="Font Size:", font=ctk.CTkFont(weight="bold")).grid(
            row=1, column=0, sticky="w", pady=(15, 5))
        self.font_size_entry = ctk.CTkEntry(form_frame, placeholder_text="Enter font size (e.g., 12)")
        self.font_size_entry.grid(row=1, column=1, sticky="ew", padx=(10, 0), pady=(15, 5))
        self.font_size_entry.insert(0, "12")
        
        # Font color
        ctk.CTkLabel(form_frame, text="Font Color:", font=ctk.CTkFont(weight="bold")).grid(
            row=2, column=0, sticky="w", pady=(15, 5))
        self.font_color_combo = ctk.CTkComboBox(form_frame, values=list(FONT_COLORS.keys()))
        self.font_color_combo.grid(row=2, column=1, sticky="ew", padx=(10, 0), pady=(15, 5))
        self.font_color_combo.set("Black")
        
        # Save location
        ctk.CTkLabel(form_frame, text="Save Location:", font=ctk.CTkFont(weight="bold")).grid(
            row=3, column=0, sticky="w", pady=(15, 5))
        
        location_frame = ctk.CTkFrame(form_frame, fg_color="transparent")
        location_frame.grid(row=3, column=1, sticky="ew", padx=(10, 0), pady=(15, 5))
        location_frame.grid_columnconfigure(0, weight=1)
        
        self.location_entry = ctk.CTkEntry(location_frame, placeholder_text="Select save directory")
        self.location_entry.grid(row=0, column=0, sticky="ew", padx=(0, 10))
        self.location_entry.insert(0, os.path.expanduser("~/Desktop"))
        
        browse_btn = ctk.CTkButton(
            location_frame, 
            text="Browse", 
            width=80,
            command=self.browse_location,
            fg_color=BIGIS_COLORS['secondary'],
            hover_color=BIGIS_COLORS['primary']
        )
        browse_btn.grid(row=0, column=1)
        
        # Instructions
        instructions_frame = ctk.CTkFrame(main_frame)
        instructions_frame.grid(row=2, column=0, sticky="ew", padx=20, pady=(0, 20))
        
        instructions_label = ctk.CTkLabel(
            instructions_frame,
            text="ℹ️ Instructions:\n\n• Enter your preferred settings above\n• Click 'Proceed' to start the ranking tracker\n• The application will open Chrome for CAPTCHA solving if needed\n• Results will be saved in the specified location",
            font=ctk.CTkFont(size=12),
            justify="left",
            anchor="w"
        )
        instructions_label.grid(row=0, column=0, padx=20, pady=15)
        
        # Proceed button
        proceed_btn = ctk.CTkButton(
            main_frame,
            text="🚀 Proceed to Tracking",
            font=ctk.CTkFont(size=16, weight="bold"),
            height=50,
            command=self.proceed,
            fg_color=BIGIS_COLORS['accent'],
            hover_color=BIGIS_COLORS['danger']
        )
        proceed_btn.grid(row=3, column=0, sticky="ew", padx=20, pady=(0, 20))
        
        # Center window
        self.center_window()
    
    def center_window(self):
        """Center the window on screen"""
        self.window.update_idletasks()
        width = self.window.winfo_width()
        height = self.window.winfo_height()
        x = (self.window.winfo_screenwidth() // 2) - (width // 2)
        y = (self.window.winfo_screenheight() // 2) - (height // 2)
        self.window.geometry(f"{width}x{height}+{x}+{y}")
    
    def browse_location(self):
        """Browse for save location"""
        folder = filedialog.askdirectory(
            title="Select Save Location",
            initialdir=self.location_entry.get() or os.path.expanduser("~")
        )
        if folder:
            self.location_entry.delete(0, ctk.END)
            self.location_entry.insert(0, folder)
    
    def validate_inputs(self):
        """Validate all inputs"""
        filename = self.filename_entry.get().strip()
        if not filename:
            messagebox.showerror("Error", "Please enter a filename")
            return False
        
        try:
            font_size = int(self.font_size_entry.get().strip())
            if font_size < 8 or font_size > 72:
                messagebox.showerror("Error", "Font size must be between 8 and 72")
                return False
        except ValueError:
            messagebox.showerror("Error", "Please enter a valid font size (number)")
            return False
        
        font_color = self.font_color_combo.get()
        if font_color not in FONT_COLORS:
            messagebox.showerror("Error", "Please select a valid font color")
            return False
        
        save_location = self.location_entry.get().strip()
        if not save_location or not os.path.exists(save_location):
            messagebox.showerror("Error", "Please select a valid save location")
            return False
        
        return True
    
    def proceed(self):
        """Proceed to tracking window"""
        if not self.validate_inputs():
            return
        
        self.config_data = {
            'filename': self.filename_entry.get().strip(),
            'font_size': int(self.font_size_entry.get().strip()),
            'font_color': FONT_COLORS[self.font_color_combo.get()],
            'save_location': self.location_entry.get().strip()
        }
        
        self.window.destroy()
        self.on_complete_callback(self.config_data)
    
    def on_closing(self):
        """Handle window closing"""
        self.window.quit()
        self.window.destroy()
        sys.exit(0)
    
    def run(self):
        """Run the configuration window"""
        try:
            self.window.mainloop()
        except KeyboardInterrupt:
            self.on_closing()
        except Exception as e:
            logging.error(f"Configuration window error: {str(e)}")
            self.on_closing()

# ==================== MAIN TRACKING WINDOW ====================

class TrackingWindow:
    """Main tracking window for BART"""
    
    def __init__(self, config):
        self.config = config
        self.window = None
        self.is_tracking = False
        self.tracker = None
        
        self.create_window()
    
    def create_window(self):
        """Create the main tracking window"""
        self.window = ctk.CTk()
        self.window.title("BART - Bigis Automated Rank Tracer")
        self.window.geometry("1000x700")
        self.window.resizable(True, True)
        
        # Handle window close event
        self.window.protocol("WM_DELETE_WINDOW", self.on_closing)
        
        # Configure grid
        self.window.grid_columnconfigure(0, weight=1)
        self.window.grid_rowconfigure(1, weight=1)
        
        # Header
        header_frame = ctk.CTkFrame(self.window, height=80, fg_color=BIGIS_COLORS['primary'])
        header_frame.grid(row=0, column=0, sticky="ew", padx=0, pady=0)
        header_frame.grid_columnconfigure(1, weight=1)
        header_frame.grid_propagate(False)
        
        # Logo
        logo_label = ctk.CTkLabel(
            header_frame,
            text="📊",
            font=ctk.CTkFont(size=32),
            text_color=BIGIS_COLORS['accent']
        )
        logo_label.grid(row=0, column=0, padx=(20, 10), pady=15)
        
        # Title and subtitle
        title_frame = ctk.CTkFrame(header_frame, fg_color="transparent")
        title_frame.grid(row=0, column=1, sticky="w", pady=15)
        
        title_label = ctk.CTkLabel(
            title_frame,
            text="BART - Bigis Automated Rank Tracer",
            font=ctk.CTkFont(size=24, weight="bold"),
            text_color=BIGIS_COLORS['white']
        )
        title_label.grid(row=0, column=0, sticky="w")
        
        subtitle_label = ctk.CTkLabel(
            title_frame,
            text="Powered by Bigis Technology",
            font=ctk.CTkFont(size=12),
            text_color=BIGIS_COLORS['light']
        )
        subtitle_label.grid(row=1, column=0, sticky="w")
        
        # Status indicator
        self.status_label = ctk.CTkLabel(
            header_frame,
            text="Ready",
            font=ctk.CTkFont(size=14, weight="bold"),
            text_color=BIGIS_COLORS['accent']
        )
        self.status_label.grid(row=0, column=2, padx=(10, 20), pady=15)
        
        # Main content area
        main_frame = ctk.CTkFrame(self.window)
        main_frame.grid(row=1, column=0, sticky="nsew", padx=20, pady=(10, 20))
        main_frame.grid_columnconfigure(1, weight=2)
        main_frame.grid_rowconfigure(0, weight=1)
        
        # Left panel - Input controls
        left_panel = ctk.CTkFrame(main_frame)
        left_panel.grid(row=0, column=0, sticky="nsew", padx=(0, 10))
        left_panel.grid_columnconfigure(0, weight=1)
        
        # Input form
        input_frame = ctk.CTkFrame(left_panel)
        input_frame.grid(row=0, column=0, sticky="ew", padx=20, pady=20)
        input_frame.grid_columnconfigure(0, weight=1)
        
        # Keyword input (now a textbox for multiple keywords)
        ctk.CTkLabel(input_frame, text="🎯 Keywords to Track (max 50, one per line or comma-separated):", font=ctk.CTkFont(weight="bold")).grid(
            row=0, column=0, sticky="w", pady=(10, 5))
        self.keyword_textbox = ctk.CTkTextbox(
            input_frame, 
            height=150,
            font=ctk.CTkFont(size=14)
        )
        self.keyword_textbox.grid(row=1, column=0, sticky="ew", pady=(0, 15))
        
        # Target domain input
        ctk.CTkLabel(input_frame, text="🌐 Target Domain:", font=ctk.CTkFont(weight="bold")).grid(
            row=2, column=0, sticky="w", pady=(0, 5))
        self.domain_entry = ctk.CTkEntry(
            input_frame, 
            placeholder_text="example.com",
            height=40,
            font=ctk.CTkFont(size=14)
        )
        self.domain_entry.grid(row=3, column=0, sticky="ew", pady=(0, 15))
        
        # Page limit
        ctk.CTkLabel(input_frame, text="📄 Page Limit:", font=ctk.CTkFont(weight="bold")).grid(
            row=4, column=0, sticky="w", pady=(0, 5))
        self.page_limit_entry = ctk.CTkEntry(
            input_frame, 
            placeholder_text="Max pages to scan",
            height=40,
            font=ctk.CTkFont(size=14)
        )
        self.page_limit_entry.grid(row=5, column=0, sticky="ew", pady=(0, 20))
        self.page_limit_entry.insert(0, "10")
        
        # Start button
        self.start_btn = ctk.CTkButton(
            input_frame,
            text="🚀 Start Tracking",
            font=ctk.CTkFont(size=16, weight="bold"),
            height=50,
            command=self.start_tracking,
            fg_color=BIGIS_COLORS['success'],
            hover_color=BIGIS_COLORS['primary']
        )
        self.start_btn.grid(row=6, column=0, sticky="ew", pady=(0, 10))
        
        # Configuration info
        config_frame = ctk.CTkFrame(left_panel)
        config_frame.grid(row=1, column=0, sticky="ew", padx=20, pady=(0, 20))
        
        ctk.CTkLabel(config_frame, text="⚙️ Configuration", font=ctk.CTkFont(weight="bold")).grid(
            row=0, column=0, sticky="w", padx=15, pady=(15, 10))
        
        config_text = f"""📁 Save Location: {self.config['save_location']}
📄 Filename: {self.config['filename']}.docx
🔤 Font Size: {self.config['font_size']}pt
🎨 Font Color: {[k for k, v in FONT_COLORS.items() if v == self.config['font_color']][0]}"""
        
        config_info = ctk.CTkLabel(
            config_frame,
            text=config_text,
            font=ctk.CTkFont(size=11),
            justify="left",
            anchor="w"
        )
        config_info.grid(row=1, column=0, sticky="w", padx=15, pady=(0, 15))
        
        # Right panel - Logs
        right_panel = ctk.CTkFrame(main_frame)
        right_panel.grid(row=0, column=1, sticky="nsew", padx=(10, 0))
        right_panel.grid_columnconfigure(0, weight=1)
        right_panel.grid_rowconfigure(1, weight=1)
        
        # Log header
        log_header = ctk.CTkFrame(right_panel, height=50, fg_color=BIGIS_COLORS['secondary'])
        log_header.grid(row=0, column=0, sticky="ew", padx=0, pady=0)
        log_header.grid_columnconfigure(0, weight=1)
        log_header.grid_propagate(False)
        
        ctk.CTkLabel(
            log_header,
            text="📋 Tracking Logs",
            font=ctk.CTkFont(size=16, weight="bold"),
            text_color=BIGIS_COLORS['white']
        ).grid(row=0, column=0, pady=15)
        
        clear_btn = ctk.CTkButton(
            log_header,
            text="Clear",
            width=80,
            height=30,
            command=self.clear_logs,
            fg_color=BIGIS_COLORS['danger'],
            hover_color=BIGIS_COLORS['warning']
        )
        clear_btn.grid(row=0, column=1, padx=(0, 15), pady=15)
        
        # Log text area
        self.log_text = ctk.CTkTextbox(
            right_panel,
            font=ctk.CTkFont(family="Consolas", size=12),
            wrap="word"
        )
        self.log_text.grid(row=1, column=0, sticky="nsew", padx=15, pady=(0, 15))
        
        # Initial welcome message
        self.log_message("🎯 Welcome to BART - Bigis Automated Rank Tracer")
        self.log_message("📊 Powered by Bigis Technology")
        self.log_message("=" * 60)
        self.log_message("📋 Instructions:")
        self.log_message("1. Enter up to 50 keywords (one per line or comma-separated)")
        self.log_message("2. Enter your target domain (without http://)")
        self.log_message("3. Set the maximum pages to scan (default: 10)")
        self.log_message("4. Click 'Start Tracking' to begin")
        self.log_message("=" * 60)
        self.log_message("⚡ Chrome will open visibly for CAPTCHA solving if needed")
        self.log_message("📄 Results will be saved to your configured Word document")
        self.log_message("")
        
        # Center window
        self.center_window()
    
    def center_window(self):
        """Center the window on screen"""
        self.window.update_idletasks()
        width = self.window.winfo_width()
        height = self.window.winfo_height()
        x = (self.window.winfo_screenwidth() // 2) - (width // 2)
        y = (self.window.winfo_screenheight() // 2) - (height // 2)
        self.window.geometry(f"{width}x{height}+{x}+{y}")
    
    def log_message(self, message):
        """Add message to log"""
        timestamp = datetime.now().strftime("%H:%M:%S")
        log_entry = f"[{timestamp}] {message}\n"
        
        # Use after to ensure thread safety
        self.window.after(0, self._update_log, log_entry)
    
    def _update_log(self, message):
        """Update log text widget (must be called from main thread)"""
        self.log_text.insert("end", message)
        self.log_text.see("end")
    
    def clear_logs(self):
        """Clear the log text area"""
        self.log_text.delete("1.0", "end")
    
    def update_status(self, status):
        """Update status label"""
        self.window.after(0, lambda: self.status_label.configure(text=status))
    
    def validate_tracking_inputs(self):
        """Validate tracking inputs"""
        # Get and process keywords
        keywords_input = self.keyword_textbox.get("1.0", "end").strip()
        if not keywords_input:
            messagebox.showerror("Error", "Please enter at least one keyword to track")
            return False
        
        # Split by commas or newlines
        keywords = [k.strip() for k in re.split(r'[,\n]', keywords_input) if k.strip()]
        if len(keywords) > 50:
            messagebox.showerror("Error", "Maximum 50 keywords allowed")
            return False
        
        domain = self.domain_entry.get().strip()
        if not domain:
            messagebox.showerror("Error", "Please enter a target domain")
            return False
        
        try:
            page_limit = int(self.page_limit_entry.get().strip())
            if page_limit < 1 or page_limit > 20:
                messagebox.showerror("Error", "Page limit must be between 1 and 20")
                return False
        except ValueError:
            messagebox.showerror("Error", "Please enter a valid page limit (number)")
            return False
        
        return keywords
    
    def start_tracking(self):
        """Start the rank tracking process"""
        if self.is_tracking:
            return
        
        keywords = self.validate_tracking_inputs()
        if not keywords:
            return
        
        self.is_tracking = True
        
        # Update UI
        self.start_btn.configure(
            text="⏳ Tracking in Progress...",
            state="disabled",
            fg_color=BIGIS_COLORS['gray']
        )
        
        # Start tracking in background thread
        thread = threading.Thread(target=self.run_tracking, args=(keywords,))
        thread.daemon = True
        thread.start()
    
    def run_tracking(self, keywords):
        """Run the tracking process for multiple keywords in background thread"""
        try:
            domain = self.domain_entry.get().strip()
            page_limit = int(self.page_limit_entry.get().strip())
            
            self.log_message(f"🚀 Starting rank tracking session for {len(keywords)} keywords...")
            self.log_message(f"🌐 Target Domain: {domain}")
            self.log_message(f"📄 Max Pages: {page_limit}")
            self.log_message("=" * 60)
            
            doc_path = None
            for idx, keyword in enumerate(keywords, 1):
                self.log_message(f"📋 Processing keyword {idx}/{len(keywords)}: '{keyword}'")
                
                # Create tracker for each keyword
                self.tracker = RankTracker(
                    keyword=keyword,
                    target_domain=domain,
                    max_pages=page_limit,
                    config=self.config,
                    log_callback=self.log_message,
                    status_callback=self.update_status
                )
                
                # Track ranking
                result = self.tracker.track_ranking()
                
                # Generate Word document
                if result:
                    self.log_message("📄 Generating Word document...")
                    try:
                        doc_path = self.tracker.create_word_document(result)
                        
                        if result['found']:
                            self.log_message(f"✅ SUCCESS! Domain found on page {result['page']}")
                            self.log_message(f"📄 Document saved: {doc_path}")
                        else:
                            self.log_message(f"❌ Domain not found for keyword: {keyword}")
                            self.log_message(f"📄 Document saved: {doc_path}")
                            
                    except Exception as e:
                        self.log_message(f"❌ Error generating document: {str(e)}")
                        self.window.after(0, lambda err=str(e): messagebox.showerror("Document Error", f"Could not generate Word document: {err}"))
            
            self.log_message("=" * 60)
            self.log_message("🎉 Tracking session completed!")
            if doc_path:
                self.window.after(0, lambda: messagebox.showinfo("Complete", f"Tracking completed for {len(keywords)} keywords.\nDocument saved: {doc_path}"))
            
        except KeyboardInterrupt:
            self.log_message("⚠️ Tracking interrupted by user")
            if self.tracker and self.tracker.driver:
                try:
                    self.tracker.driver.quit()
                except:
                    pass
        except Exception as e:
            self.log_message(f"❌ Error during tracking: {str(e)}")
            logging.error(f"Tracking error: {str(e)}")
            self.window.after(0, lambda: messagebox.showerror("Error", f"Tracking failed: {str(e)}"))
        
        finally:
            # Reset UI state
            self.window.after(0, self._reset_ui_state)
    
    def _reset_ui_state(self):
        """Reset UI state after tracking"""
        self.is_tracking = False
        self.start_btn.configure(
            text="🚀 Start Tracking",
            state="normal",
            fg_color=BIGIS_COLORS['success']
        )
        self.update_status("Ready")
    
    def on_closing(self):
        """Handle window closing"""
        if self.is_tracking and self.tracker:
            # Ask user if they want to stop tracking
            result = messagebox.askyesno("Confirm Exit", "Tracking is in progress. Do you want to stop and exit?")
            if not result:
                return
            
            # Stop tracking
            try:
                if self.tracker.driver:
                    self.tracker.driver.quit()
            except:
                pass
        
        self.window.quit()
        self.window.destroy()
        sys.exit(0)
    
    def run(self):
        """Run the tracking window"""
        try:
            self.window.mainloop()
        except KeyboardInterrupt:
            self.on_closing()
        except Exception as e:
            logging.error(f"Tracking window error: {str(e)}")
            self.on_closing()

# ==================== MAIN APPLICATION CLASS ====================

class BARTApplication:
    """Main BART Application Class"""
    
    def __init__(self):
        self.config_data = None
        self.tracking_window = None
    
    def start_application(self):
        """Start the application with configuration window"""
        config_window = ConfigurationWindow(self.on_configuration_complete)
        config_window.run()
    
    def on_configuration_complete(self, config_data):
        """Handle configuration completion and transition to tracking window"""
        self.config_data = config_data
        
        # Start tracking window
        self.tracking_window = TrackingWindow(self.config_data)
        self.tracking_window.run()

# ==================== MAIN ENTRY POINT ====================

def main():
    """Main entry point"""
    try:
        print("🚀 Starting BART - Bigis Automated Rank Tracer")
        print("📊 Professional Google Ranking Tracker by Bigis Technology")
        print("=" * 60)
        
        app = BARTApplication()
        app.start_application()
        
    except KeyboardInterrupt:
        print("\n👋 BART application terminated by user")
        sys.exit(0)
    except Exception as e:
        print(f"❌ Error starting BART: {str(e)}")
        logging.error(f"Application error: {str(e)}")
        try:
            messagebox.showerror("Error", f"Failed to start BART: {str(e)}")
        except:
            pass  # GUI may not be available
        sys.exit(1)

if __name__ == "__main__":
    main()